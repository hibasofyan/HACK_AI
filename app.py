import os
import gradio as gr
import google.generativeai as genai 
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv
import speech_recognition as sr
import time
from gtts import gTTS
import tempfile
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import inch
from rich.console import Console
from rich.panel import Panel
from rich.text import Text
from rich.table import Table as RichTable
from rich import box
import re
import json
from typing import List 
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.agents import AgentExecutor, create_tool_calling_agent
from langchain_core.prompts import ChatPromptTemplate
from langchain.tools import tool


load_dotenv()


GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=GOOGLE_API_KEY)

# Définition des outils
@tool
def read_document_tool(file_path: str) -> str:
    """Lit le contenu d'un fichier PDF ou DOCX situé à file_path et le retourne sous forme de texte."""
    print(f"Outil : Lecture du document depuis {file_path}")
    if not file_path:
        return "Erreur : Aucun chemin de fichier fourni."
    try:
        if file_path.lower().endswith('.pdf'):
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            print("Outil : PDF lu avec succès.")
            return text
        elif file_path.lower().endswith('.docx'):
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            print("Outil : DOCX lu avec succès.")
            return text
        else:
            print("Outil : Format de fichier non supporté.")
            return "Format de fichier non supporté. Veuillez fournir un fichier PDF ou DOCX."
    except Exception as e:
        print(f"Outil : Erreur lors de la lecture du fichier - {e}")
        return f"Erreur lors de la lecture du fichier : {e}"

@tool
def analyze_cv_with_llm_tool(cv_text: str, job_description: str) -> str:
    """Analyse le texte du CV fourni par rapport à la description du poste en utilisant un LLM et retourne l'analyse structurée."""
    print("Outil : Analyse du CV avec le LLM...")
    if not cv_text or not job_description:
        return "Erreur : Texte du CV ou description du poste manquant."
        
    prompt = f"""
    Analysez le CV suivant par rapport à la description du poste.
    
    Description du poste :
    {job_description}
    
    CV :
    {cv_text}
    
    Fournissez une analyse structurée en Markdown incluant :
    Nom du Candidat : [Nom et Prénom du candidat ici]
    Poste Visé : [Intitulé du poste ici]
    
    ## Analyse du Candidat
    
    ### Points forts
    Liste des points forts pertinents pour le poste.
    
    ### Points faibles
    Liste des points faibles ou des domaines nécessitant un développement.
    
    ### Score de pertinence
    Un score global sur 100 évaluant l'adéquation du candidat avec le poste.
    
    ### Résumé du profil
    Un résumé concis du profil du candidat et de son expérience.
    """
    
    try:
        response = llm.invoke(prompt)
        analysis_text = response.content if response and response.content else "Analyse non disponible."
        print("Outil : Analyse du CV par le LLM terminée.")
        return analysis_text
    except Exception as e:
        print(f"Outil : Erreur lors de l'analyse par le LLM - {e}")
        return f"Erreur lors de l'analyse par le LLM : {e}"

@tool
def analyze_sentiment_tool(text: str) -> str:
    """Analyse le sentiment du texte donné dans le contexte d'un entretien d'embauche (ex: positif, négatif, neutre, hésitant, confiant, stressé, enthousiaste) et retourne un terme unique ou une courte phrase décrivant le sentiment."""
    print("Outil : Analyse du sentiment...")
    if not text or text in ["Je n'ai pas compris votre réponse", "Erreur de service de reconnaissance vocale", "Je n'ai pas détecté de parole", "Erreur lors de l'enregistrement vocal"]:
        return "Non analysable"

    prompt = f"""
    Analysez le sentiment principal de la réponse suivante du candidat dans le contexte d'un entretien d'embauche. Est-il positif, négatif, neutre, hésitant, confiant, stressé, enthousiaste ? Choisissez le terme qui décrit le mieux le sentiment général.
    
    Réponse du candidat : {text}
    
    Répondez uniquement avec un terme unique ou une courte phrase décrivant le sentiment.
    """
    try:
        response = llm.invoke(prompt)
        sentiment = response.content.strip() if response and response.content else "Analyse impossible"
        print(f"Outil : Analyse du sentiment terminée : {sentiment}")
        return sentiment
    except Exception as e:
        print(f"Outil : Erreur lors de l'analyse du sentiment - {e}")
        return f"Erreur lors de l'analyse du sentiment : {e}"

@tool
def generate_next_interview_question_tool(cv_analysis: str, interview_transcript: List[dict], interview_goals: List[str], achieved_goals: List[str]) -> str:
    """Génère la prochaine question d'entretien pertinente et ouverte basée sur l'analyse du CV, la transcription de l'entretien, les objectifs et les objectifs déjà atteints."""
    print("Outil : Génération de la prochaine question d'entretien...")
    
    prompt = f"""
    Vous êtes un recruteur professionnel menant un entretien d'embauche.
    Sur la base de l'analyse du CV, de la transcription de l'entretien, des objectifs et des objectifs déjà atteints :
    
    Analyse du CV : {cv_analysis}
    
    Objectifs de l'entretien : {interview_goals}
    Objectifs atteints : {achieved_goals}
    Questions et réponses précédentes :
    {interview_transcript}
    
    En considérant les objectifs qui n'ont PAS ENCORE été suffisamment couverts, formulez une question d'entretien pertinente et ouverte pour continuer à évaluer le candidat.
    La question doit être directe, professionnelle et s'appuyer sur la conversation précédente.
    IMPORTANT : Répondez UNIQUEMENT avec la question, sans explications ni commentaires.
    """
    
    try:
        response = llm.invoke(prompt)
        question = response.content.strip() if response and response.content else "Pouvez-vous me parler de votre expérience professionnelle la plus significative ?"
        print(f"Outil : Prochaine question d'entretien générée : '{question}'")
        return question
    except Exception as e:
        print(f"Outil : Erreur lors de la génération de la question - {e}")
        return f"Erreur lors de la génération de la question : {e}"

@tool
def evaluate_interview_progress_tool(interview_transcript: List[dict], interview_goals: List[str], achieved_goals: List[str]) -> str:
    """Évalue la progression de l'entretien sur la base de la transcription, des objectifs et des objectifs atteints. Retourne 'CONTINUER' si l'entretien doit continuer pour couvrir les objectifs restants, ou 'FIN' si les objectifs sont suffisamment couverts ou si le nombre maximum de questions est atteint."""
    print("Outil : Évaluation de la progression de l'entretien...")

    if len(achieved_goals) >= len(interview_goals) and len(interview_goals) > 0:
        print("Outil : Tous les objectifs semblent atteints. Recommandation : FIN.")
        return "FIN"

    prompt = f"""
    Sur la base de la transcription de l'entretien, des objectifs et des objectifs atteints :

    Objectifs de l'entretien : {interview_goals}
    Objectifs atteints : {achieved_goals}
    Transcription de l'entretien :
    {interview_transcript}

    Compte tenu de la progression et des informations obtenues, est-il judicieux de continuer l'entretien pour mieux évaluer le candidat sur les objectifs non encore couverts ?
    Répondez UNIQUEMENT avec 'CONTINUER' ou 'FIN'.
    """

    try:
        response = llm.invoke(prompt)
        decision = response.content.strip().upper() if response and response.content else "CONTINUER"
        print(f"Outil : Résultat de l'évaluation de la progression : {decision}")
        return decision
    except Exception as e:
        print(f"Outil : Erreur lors de l'évaluation de la progression - {e}")
        return "CONTINUER" # Continuer par défaut en cas d'erreur

@tool
def generate_final_report_content_tool(cv_analysis: str, interview_transcript: List[dict], interview_goals: List[str]) -> str:
    """Analyse l'analyse du CV, la transcription de l'entretien et les objectifs pour générer le contenu complet du rapport de recrutement final. Retourne une chaîne JSON contenant les sections structurées du rapport."""
    print("Outil : Génération du contenu du rapport final...")

    if not cv_analysis or not interview_transcript or not interview_goals:
        return "Erreur : Données manquantes (analyse du CV, transcription ou objectifs) pour générer le rapport final."

    prompt = f"""
    En tant qu'expert en recrutement, analysez l'entretien et générez un rapport détaillé au format JSON.
    Basez-vous sur l'analyse du CV, la transcription de l'entretien et les objectifs définis.

    Analyse du CV :
    {cv_analysis}

    Objectifs de l'entretien :
    {interview_goals}

    Transcription de l'entretien :
    {interview_transcript}

    Générez un rapport JSON structuré avec les sections suivantes :
    {{
        "overall_evaluation": "Une évaluation globale détaillée de la performance du candidat pendant l'entretien, en se basant sur ses réponses et son comportement",
        "demonstrated_strengths": [
            "Liste des points forts démontrés pendant l'entretien, avec des exemples concrets tirés des réponses"
        ],
        "areas_for_improvement": [
            "Liste des points à améliorer identifiés pendant l'entretien, avec des exemples concrets"
        ],
        "interview_score": "Un score sur 100 avec justification détaillée",
        "final_recommendation": "Une recommandation claire (ex: 'Recommandé pour un second entretien', 'Refusé', 'À considérer pour un autre poste') avec justification",
        "goals_achievement": {{
            "achieved_goals": [
                "Liste des objectifs qui ont été suffisamment évalués pendant l'entretien"
            ],
            "unachieved_goals": [
                "Liste des objectifs qui n'ont pas été suffisamment couverts"
            ]
        }}
    }}

    IMPORTANT :
    - Basez votre analyse uniquement sur les réponses données pendant l'entretien
    - Soyez spécifique et citez des exemples concrets des réponses du candidat
    - Évitez les généralités et les clichés
    - Assurez-vous que chaque section est détaillée et justifiée
    - Le format de sortie doit être un JSON valide
    """

    try:
        response = llm.invoke(prompt)
        report_content = response.content.strip() if response and response.content else "Impossible de générer le contenu du rapport final."
        print("Outil : Contenu du rapport final généré par le LLM.")
        
        # Nettoyer la réponse pour s'assurer qu'elle est en JSON valide
        report_content = re.sub(r'^```json\n', '', report_content)
        report_content = re.sub(r'\n```$', '', report_content)
        report_content = report_content.strip()
        
        try:
            json.loads(report_content)
            return report_content
        except json.JSONDecodeError:
            default_report = {
                "overall_evaluation": "Évaluation non disponible - Format de réponse invalide",
                "demonstrated_strengths": ["Données insuffisantes"],
                "areas_for_improvement": ["Données insuffisantes"],
                "interview_score": "Non évalué",
                "final_recommendation": "Recommandation non disponible",
                "goals_achievement": {
                    "achieved_goals": [],
                    "unachieved_goals": interview_goals
                }
            }
            return json.dumps(default_report, ensure_ascii=False)
            
    except Exception as e:
        print(f"Outil : Erreur lors de la génération du contenu du rapport final - {e}")
        return json.dumps({
            "overall_evaluation": f"Erreur lors de la génération du rapport : {str(e)}",
            "demonstrated_strengths": ["Données insuffisantes"],
            "areas_for_improvement": ["Données insuffisantes"],
            "interview_score": "Non évalué",
            "final_recommendation": "Recommandation non disponible",
            "goals_achievement": {
                "achieved_goals": [],
                "unachieved_goals": interview_goals
            }
        }, ensure_ascii=False)

@tool
def determine_interview_goals_tool(cv_analysis: str, job_description: str) -> List[str]:
    """Détermine les objectifs clés de l'entretien basés sur l'analyse du CV et la description du poste. Retourne une liste d'objectifs d'évaluation spécifiques."""
    print("Outil : Détermination des objectifs de l'entretien...")
    if not cv_analysis or not job_description:
        return ["Erreur : Analyse du CV ou description du poste manquante pour déterminer les objectifs."]
        
    prompt = f"""
    Sur la base de l'analyse du CV et de la description du poste fournies, identifiez 3 à 5 objectifs clés que l'entretien doit viser pour évaluer correctement le candidat.
    Ces objectifs doivent être spécifiques et axés sur l'évaluation des compétences et de l'adéquation au poste.
    Formatez votre réponse comme une liste JSON de chaînes, où chaque chaîne est un objectif.
    Exemple : ["Évaluer l'expérience en gestion de projet", "Vérifier les compétences techniques en Python", "Comprendre la motivation pour le poste"]

    Analyse du CV :
    {cv_analysis}

    Description du poste :
    {job_description}
    """
    
    try:
        response = llm.invoke(prompt)
        goals_raw = response.content.strip() if response and response.content else "[]"
        print(f"Outil : Réponse brute des objectifs : {goals_raw}")
        
        goals_cleaned = re.sub(r'^```json\n', '', goals_raw)
        goals_cleaned = re.sub(r'\n```$', '', goals_cleaned)
        goals_cleaned = goals_cleaned.strip()
        print(f"Outil : Réponse nettoyée des objectifs : {goals_cleaned}")

        try:
            goals_list = json.loads(goals_cleaned)
            if not isinstance(goals_list, list):
                 print("Outil : Attention - Le LLM n'a pas retourné une liste JSON après nettoyage.")
                 return ["Impossible de parser les objectifs de la réponse du LLM (pas une liste)."]
            return goals_list if goals_list else ["Aucun objectif spécifique déterminé."]
        except json.JSONDecodeError:
             print("Outil : Erreur lors du décodage du JSON des objectifs après nettoyage.")
             return ["Erreur lors du décodage des objectifs de la réponse du LLM. Réponse : '" + goals_cleaned[:100] + "...'"]
             
    except Exception as e:
        print(f"Outil : Erreur lors de la détermination des objectifs - {e}")
        return [f"Erreur lors de la détermination des objectifs : {e}"]

# Liste de tous les outils disponibles
tools = [
    read_document_tool,
    analyze_cv_with_llm_tool,
    analyze_sentiment_tool,
    generate_next_interview_question_tool,
    evaluate_interview_progress_tool,
    generate_final_report_content_tool,
    determine_interview_goals_tool,
]


# Définition de l'agent Analyseur de CV
def create_cv_analyzer_agent(llm, tools):
    prompt = ChatPromptTemplate.from_messages([
        ("system", "Vous êtes un expert en analyse de CV. Votre rôle est d'analyser le CV d'un candidat par rapport à une description de poste et de fournir un rapport structuré."),
        ("human", "Analysez le CV suivant en fonction de la description du poste. Utilisez les outils disponibles si nécessaire.\nDescription du poste: {job_description}\nTexte du CV: {cv_text}"),
        ("placeholder", "{agent_scratchpad}"),
    ])
    
    agent = create_tool_calling_agent(llm, tools, prompt)
    agent_executor = AgentExecutor(agent=agent, tools=tools, verbose=True)
    
    return agent_executor

# Définition de l'agent Simulateur d'Entretien
def create_interview_agent(llm, tools):
    prompt = ChatPromptTemplate.from_messages([
        ("system", """Vous êtes un recruteur professionnel IA. Votre rôle est de mener un entretien structuré basé sur l'analyse du CV du candidat et les objectifs définis.
        IMPORTANT: Vous devez TOUJOURS répondre avec soit :
        1. Une question d'entretien claire et professionnelle qui s'appuie sur la conversation précédente et aide à évaluer les objectifs restants
        2. Le mot "FIN" si vous déterminez que l'entretien doit se conclure
        
        Ne répondez jamais avec une chaîne vide ou un autre format."""),
        ("human", """
État actuel de l'entretien :
Analyse du CV : {cv_analysis}
Objectifs de l'entretien : {interview_goals}
Transcription de l'entretien : {interview_transcript}
Dernière réponse du candidat : {candidate_response}

Sur la base de ces informations, veuillez générer la prochaine question d'entretien ou déterminer si l'entretien doit se terminer.
Rappelez-vous : Vous devez répondre soit avec une question claire, soit avec "FIN".
"""),
         ("placeholder", "{agent_scratchpad}"), # Pour le processus de réflexion de l'agent
    ])
    
    # Filtrer les outils pour n'inclure que ceux pertinents pour l'entretien
    interview_tools = [tool for tool in tools if tool.name in ["generate_next_interview_question_tool", "analyze_sentiment_tool", "evaluate_interview_progress_tool"]]

    agent = create_tool_calling_agent(llm, interview_tools, prompt)
    agent_executor = AgentExecutor(agent=agent, tools=interview_tools, verbose=True)
    
    return agent_executor

# Définition de l'agent Analyste Post-Entretien
def create_post_interview_analyzer_agent(llm, tools):
    prompt = ChatPromptTemplate.from_messages([
        ("system", "Vous êtes un expert en analyse post-entretien. Votre rôle est d'analyser l'analyse du CV, la transcription de l'entretien et les objectifs pour générer un rapport de recrutement final complet."),
        ("human", "Générez le rapport de recrutement final sur la base du contexte suivant.\nAnalyse du CV : {cv_analysis}\nObjectifs de l'entretien : {interview_goals}\nTranscription de l'entretien : {interview_transcript}"),
        ("placeholder", "{agent_scratchpad}"), 
    ])
    
    # Filtrer les outils pour n'inclure que ceux pertinents pour l'analyse post-entretien
    post_interview_tools = [tool for tool in tools if tool.name in ["generate_final_report_content_tool"]]

    agent = create_tool_calling_agent(llm, post_interview_tools, prompt)
    agent_executor = AgentExecutor(agent=agent, tools=post_interview_tools, verbose=True)
    
    return agent_executor

class SecretaireAgent:
    def __init__(self):
        self.console = Console()
        self.styles = getSampleStyleSheet()
        self._setup_styles()

    def _setup_styles(self):
        """Configure les styles personnalisés pour le PDF"""
        # Style pour le titre principal
        self.styles.add(ParagraphStyle(
            name='CVMainTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#1A237E'),
            alignment=1,  # Centré
            fontName='Helvetica-Bold'
        ))

        # Style pour les sous-titres principaux
        self.styles.add(ParagraphStyle(
            name='CVSectionTitle',
            parent=self.styles['Heading2'],
            fontSize=18,
            spaceAfter=20,
            textColor=colors.HexColor('#283593'),
            fontName='Helvetica-Bold',
            leftIndent=20
        ))

        # Style pour les sous-sections
        self.styles.add(ParagraphStyle(
            name='CVSubSection',
            parent=self.styles['Heading3'],
            fontSize=14,
            spaceAfter=15,
            textColor=colors.HexColor('#303F9F'),
            fontName='Helvetica-Bold',
            leftIndent=40
        ))

        # Style pour le texte normal
        self.styles.add(ParagraphStyle(
            name='CVNormal',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            textColor=colors.HexColor('#212121'),
            fontName='Helvetica',
            leftIndent=40,
            rightIndent=40
        ))

        # Style pour les listes à puces
        self.styles.add(ParagraphStyle(
            name='CVBullet',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            textColor=colors.HexColor('#212121'),
            fontName='Helvetica',
            leftIndent=60,
            rightIndent=40,
            bulletIndent=40
        ))

        # Style pour les informations de candidat
        self.styles.add(ParagraphStyle(
            name='CVCandidateInfo',
            parent=self.styles['Normal'],
            fontSize=12,
            spaceAfter=15,
            textColor=colors.HexColor('#424242'),
            fontName='Helvetica-Bold',
            alignment=1  # Centré
        ))

        # Nouveaux styles pour le rapport final
        self.styles.add(ParagraphStyle(
            name='ReportTitle',
            parent=self.styles['Title'],
            fontSize=28,
            spaceAfter=40,
            textColor=colors.HexColor('#0D47A1'), # Bleu plus foncé
            alignment=1, # Centré
            fontName='Helvetica-Bold'
        ))

        self.styles.add(ParagraphStyle(
            name='ReportSection',
            parent=self.styles['Heading1'],
            fontSize=20,
            spaceAfter=25,
            spaceBefore=20,
            textColor=colors.HexColor('#1565C0'), # Bleu moyen
            fontName='Helvetica-Bold',
            leftIndent=0
        ))

        self.styles.add(ParagraphStyle(
            name='ReportSubSection',
            parent=self.styles['Heading2'],
            fontSize=16,
            spaceAfter=15,
            spaceBefore=15,
            textColor=colors.HexColor('#1976D2'), # Bleu un peu plus clair
            fontName='Helvetica-Bold',
            leftIndent=20
        ))

        self.styles.add(ParagraphStyle(
            name='ReportBody',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=8,
            textColor=colors.HexColor('#212121'),
            fontName='Helvetica',
            leftIndent=20,
            rightIndent=20
        ))

        self.styles.add(ParagraphStyle(
            name='ReportQuestion',
            parent=self.styles['ReportBody'],
            fontName='Helvetica-BoldOblique',
            textColor=colors.HexColor('#303F9F'), # Indigo
        ))

        self.styles.add(ParagraphStyle(
            name='ReportAnswer',
            parent=self.styles['ReportBody'],
            spaceAfter=15,
            textColor=colors.HexColor('#424242'), # Gris foncé
        ))

        self.styles.add(ParagraphStyle(
            name='ReportSentiment',
            parent=self.styles['ReportBody'],
            fontName='Helvetica-Oblique',
            textColor=colors.HexColor('#546E7A'), # Gris bleuâtre
            spaceBefore=5,
            spaceAfter=15,
        ))

    def create_pdf_report(self, cv_analysis, output_path):
        """Crée un rapport PDF professionnel formaté selon le template."""
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        story = []
        styles = self.styles # Utiliser les styles configurés

        # --- Titre Principal ---
        story.append(Paragraph("Rapport d'Analyse de Candidature", styles['CVMainTitle']))
        story.append(Spacer(1, 30))

        # --- Table d'Informations ---
        # Extraire Nom du Candidat et Poste Visé du début de l'analyse
        candidat_match = re.search(r'Nom du Candidat : (.*?)\n', cv_analysis)
        nom_candidat = candidat_match.group(1).strip() if candidat_match else "Non spécifié"

        poste_match = re.search(r'Poste Visé : (.*?)\n', cv_analysis)
        nom_poste = poste_match.group(1).strip() if poste_match else "Non spécifié"

        # Obtenir la date du jour
        from datetime import date
        date_analyse = date.today().strftime("%d/%m/%Y")

        data = [
            ['Nom du Candidat :', nom_candidat],
            ['Poste Visé :', nom_poste],
            ["Date de l'analyse :", date_analyse],
            ['Analysé par :', 'AIRHTech'],
        ]
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BOX', (0, 0), (-1, -1), 1, colors.black),
        ])
        info_table = Table(data)
        info_table.setStyle(table_style)
        story.append(info_table)
        story.append(Spacer(1, 20))

        # --- Résumé du Profil ---
        story.append(Paragraph("Résumé du Profil", styles['CVSectionTitle']))
        story.append(Spacer(1, 10))
        resume_match = re.search(r'### Résumé du profil\n(.*?)(?:\n###|$)', cv_analysis, re.DOTALL)
        resume_text = resume_match.group(1).strip() if resume_match else "Résumé non trouvé."
        story.append(Paragraph(resume_text, styles['CVNormal']))
        story.append(Spacer(1, 20))

        # --- Points Forts ---
        story.append(Paragraph("Points Forts", styles['CVSectionTitle']))
        story.append(Spacer(1, 10))
        forts_match = re.search(r'### Points forts\n(.*?)(?:\n###|$)', cv_analysis, re.DOTALL)
        forts_text = forts_match.group(1).strip() if forts_match else "Points forts non trouvés."
        # Ajouter les points forts comme une liste à puces avec gras après la puce
        for point in forts_text.split('\n'):
            if point.strip():
                 # Retirer le marqueur de liste Markdown (*)
                 clean_point = point.strip().lstrip('* ').strip()
                 # Retirer les marqueurs Markdown de gras (**) du début et de la fin
                 clean_point = clean_point.lstrip('**').rstrip('**').strip()
                 # Séparer le texte en une partie à mettre en gras et le reste
                 parts = clean_point.split(':', 1)
                 if len(parts) > 1:
                     bold_part = parts[0].strip() + ":"
                     rest_part = parts[1].strip()
                     # Utiliser des balises HTML pour le gras dans le Paragraph
                     story.append(Paragraph(f"• <b>{bold_part}</b> {rest_part}", styles['CVBullet']))
                 else:
                     # Si pas de :, juste mettre le point en gras
                      story.append(Paragraph(f"• <b>{clean_point}</b>", styles['CVBullet']))

        story.append(Spacer(1, 20))

        # --- Points Faibles ---
        story.append(Paragraph("Points Faibles", styles['CVSectionTitle']))
        story.append(Spacer(1, 10))
        faibles_match = re.search(r'### Points faibles\n(.*?)(?:\n###|$)', cv_analysis, re.DOTALL)
        faibles_text = faibles_match.group(1).strip() if faibles_match else "Points faibles non trouvés."
         # Ajouter les points faibles comme une liste à puces avec gras après la puce
        for point in faibles_text.split('\n'):
            if point.strip():
                 # Retirer le marqueur de liste Markdown (*)
                 clean_point = point.strip().lstrip('* ').strip()
                 # Retirer les marqueurs Markdown de gras (**) du début et de la fin
                 clean_point = clean_point.lstrip('**').rstrip('**').strip()
                 # Séparer le texte en une partie à mettre en gras et le reste
                 parts = clean_point.split(':', 1)
                 if len(parts) > 1:
                     bold_part = parts[0].strip() + ":"
                     rest_part = parts[1].strip()
                     # Utiliser des balises HTML pour le gras dans le Paragraph
                     story.append(Paragraph(f"• <b>{bold_part}</b> {rest_part}", styles['CVBullet']))
                 else:
                      # Si pas de :, juste mettre le point en gras
                      story.append(Paragraph(f"• <b>{clean_point}</b>", styles['CVBullet']))

        story.append(Spacer(1, 20))

        # --- Score de Pertinence ---
        story.append(Paragraph("Score de Pertinence", styles['CVSectionTitle']))
        story.append(Spacer(1, 10))
        score_match = re.search(r'### Score de pertinence\n(.*?)(?:\n###|$)', cv_analysis, re.DOTALL)
        score_content = score_match.group(1).strip() if score_match else "Score non trouvé."
        
        score_lines = score_content.split('\n', 1)
        score_value = score_lines[0].strip() if score_lines else "N/A"
        score_justification = score_lines[1].strip() if len(score_lines) > 1 else "Justification non disponible."

        story.append(Paragraph(score_value, styles['CVSubSection']))
        story.append(Spacer(1, 5))
        story.append(Paragraph(score_justification, styles['CVNormal']))
        story.append(Spacer(1, 20))

        # Génération du PDF
        doc.build(story)

    def display_analysis(self, cv_analysis):
        """Affiche l'analyse de manière élégante dans la console"""
        # Création d'une table pour l'affichage
        table = RichTable(box=box.ROUNDED, show_header=False, padding=(0, 1))
        
        # Traitement de l'analyse
        sections = cv_analysis.split('##')
        for section in sections:
            if not section.strip():
                continue
                
            lines = section.strip().split('\n')
            title = lines[0].strip()
            content = '\n'.join(lines[1:]).strip()
            
            # Ajout du titre
            table.add_row(
                Text(title, style="bold cyan"),
                style="on rgb(40,44,52)"
            )
            
            # Traitement du contenu
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    if para.strip().startswith('* '):
                        # Traitement des listes à puces
                        items = para.split('\n')
                        for item in items:
                            if item.strip():
                                clean_item = item.strip().lstrip('* ').strip()
                                table.add_row(
                                    Text(f"• {clean_item}", style="white"),
                                    style="on rgb(40,44,52)"
                                )
                    else:
                        table.add_row(
                            Text(para.strip(), style="white"),
                            style="on rgb(40,44,52)"
                        )
        
        # Affichage dans un panel
        self.console.print(Panel(table, title="Analyse de CV", border_style="cyan"))

    # Nouvelle méthode pour créer le rapport final PDF
    def create_final_report_pdf(self, cv_analysis_text, interview_transcript_list, evaluation_json):
        """Crée le rapport final PDF avec l'analyse CV, la transcription et l'évaluation."""
        try:
            # Assurer que les entrées sont du bon type
            cv_analysis_text = str(cv_analysis_text) if cv_analysis_text is not None else ""
            if not isinstance(interview_transcript_list, list):
                print(f"Avertissement : interview_transcript_list n'est pas une liste (type: {type(interview_transcript_list)}). Initialisation à une liste vide.")
                interview_transcript_list = []
            evaluation_json = str(evaluation_json) if evaluation_json is not None else "{}"

            # Créer un fichier temporaire pour le PDF
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            temp_file.close()
            pdf_path = temp_file.name

            doc = SimpleDocTemplate(
                pdf_path,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )

            story = []
            styles = self.styles

            # --- Titre Principal du Rapport ---
            story.append(Paragraph("Rapport de Recrutement Final", styles['ReportTitle']))
            story.append(Spacer(1, 40))

            # --- Résumé de l'Analyse de CV ---
            story.append(Paragraph("1. Analyse de CV - Résumé", styles['ReportSection']))
            story.append(Spacer(1, 10))
            
            # Extraire le résumé du profil de l'analyse CV
            resume_match = re.search(r'### Résumé du profil\n(.*?)(?:\n###|$)', cv_analysis_text, re.DOTALL)
            resume_text = resume_match.group(1).strip() if resume_match else "Résumé non trouvé."
            story.append(Paragraph(resume_text, styles['ReportBody']))
            story.append(Spacer(1, 20))

            # --- Transcription de l'Entretien ---
            story.append(Paragraph("2. Transcription de l'Entretien", styles['ReportSection']))
            story.append(Spacer(1, 10))

            if not interview_transcript_list:
                story.append(Paragraph("Aucune transcription disponible.", styles['ReportBody']))
            else:
                for i, turn in enumerate(interview_transcript_list):
                    # Assurer que les éléments du tour sont des dictionnaires et que les valeurs sont des chaînes
                    if isinstance(turn, dict):
                        question = str(turn.get('question', 'Question non disponible'))
                        reponse = str(turn.get('reponse', 'Réponse non disponible'))
                        sentiment = str(turn.get('sentiment', 'Sentiment non analysé'))

                        story.append(Paragraph(f"**Tour {i+1}**", styles['ReportSubSection']))
                        story.append(Paragraph(f"**Recruteur :** {question}", styles['ReportQuestion']))
                        story.append(Paragraph(f"**Candidat :** {reponse}", styles['ReportAnswer']))
                        story.append(Paragraph(f"_Sentiment : {sentiment}_", styles['ReportSentiment']))
                    else:
                        print(f"Avertissement : Élément inattendu dans interview_transcript_list (type: {type(turn)}). Ignoré.")

            story.append(Spacer(1, 20))

            # --- Évaluation de l'Entretien ---
            story.append(Paragraph("3. Évaluation de l'Entretien", styles['ReportSection']))
            story.append(Spacer(1, 10))
            
            try:
                # Parser le JSON d'évaluation. Utilisez evaluation_json qui est déjà converti en str.
                evaluation_data = json.loads(evaluation_json)
                
                # Évaluation Globale
                story.append(Paragraph("Évaluation Globale", styles['ReportSubSection']))
                story.append(Paragraph(str(evaluation_data.get('overall_evaluation', 'Évaluation non disponible')), styles['ReportBody']))
                story.append(Spacer(1, 10))

                # Points Forts
                story.append(Paragraph("Points Forts Démontrés", styles['ReportSubSection']))
                for strength in evaluation_data.get('demonstrated_strengths', []):
                     story.append(Paragraph(f"• {str(strength)}", styles['ReportBody'])) # Convertir en str
                story.append(Spacer(1, 10))

                # Points à Améliorer
                story.append(Paragraph("Points à Améliorer", styles['ReportSubSection']))
                for area in evaluation_data.get('areas_for_improvement', []):
                     story.append(Paragraph(f"• {str(area)}", styles['ReportBody'])) # Convertir en str
                story.append(Spacer(1, 10))

                # Score
                story.append(Paragraph("Score de l'Entretien", styles['ReportSubSection']))
                story.append(Paragraph(str(evaluation_data.get('interview_score', 'Score non disponible')), styles['ReportBody']))
                story.append(Spacer(1, 10))

                # Objectifs Atteints/Non Atteints
                story.append(Paragraph("Objectifs de l'Entretien", styles['ReportSubSection']))
                goals_achievement = evaluation_data.get('goals_achievement', {})
                
                story.append(Paragraph("Objectifs Atteints :", styles['ReportBody']))
                for goal in goals_achievement.get('achieved_goals', []):
                     story.append(Paragraph(f"✓ {str(goal)}", styles['ReportBody'])) # Convertir en str
                
                story.append(Paragraph("Objectifs Non Atteints :", styles['ReportBody']))
                for goal in goals_achievement.get('unachieved_goals', []):
                     story.append(Paragraph(f"✗ {str(goal)}", styles['ReportBody'])) # Convertir en str
                story.append(Spacer(1, 10))

                # Recommandation Finale
                story.append(Paragraph("Recommandation Finale", styles['ReportSubSection']))
                story.append(Paragraph(str(evaluation_data.get('final_recommendation', 'Recommandation non disponible')), styles['ReportBody']))

            except json.JSONDecodeError as e:
                print(f"Erreur lors du parsing du JSON d'évaluation dans create_final_report_pdf : {e}")
                story.append(Paragraph(f"Erreur lors de l'analyse de l'évaluation : {e}.", styles['ReportBody']))
            except Exception as e:
                print(f"Erreur inattendue lors du traitement de l'évaluation JSON dans create_final_report_pdf : {e}")
                story.append(Paragraph(f"Erreur inattendue lors du traitement de l'évaluation : {e}.", styles['ReportBody']))

            # Génération du PDF
            doc.build(story)
            print(f"PDF final généré avec succès : {pdf_path}")
            return pdf_path

        except Exception as e:
            print(f"Erreur inattendue lors de la création du rapport PDF : {e}")
            return None

class RecrutementApp:
    def __init__(self, llm, tools):
        self.model = genai.GenerativeModel('gemini-1.5-flash') # Modèle utilisé pour certaines tâches comme la correction de transcription
        self.llm = llm # Le LLM configuré avec LangChain pour les agents
        self.tools = tools # Les outils disponibles pour les agents
        self.cv_text = ""
        self.job_description = ""
        self.interview_transcript = [] # Stocke la transcription de l'entretien [{"role": ..., "content": ...}, ...]
        self.cv_analysis = None
        self.recognizer = sr.Recognizer()
        self.current_question = None
        self.interview_started = False
        self.question_count = 0
        self.interview_goals = [] # Objectifs déterminés pour l'entretien
        self.achieved_goals = set() # Objectifs jugés suffisamment couverts
        self.conversation_history = [] # Historique complet de la conversation pour le chatbot Gradio
        self.secretaire = SecretaireAgent()
        # Créer l'agent Analyseur de CV
        self.cv_analyzer_agent = create_cv_analyzer_agent(self.llm, self.tools)
        # Créer l'agent Interviewer
        self.interview_agent = create_interview_agent(self.llm, self.tools)
        # Créer l'agent Analyste Post-Entretien
        self.post_interview_analyzer_agent = create_post_interview_analyzer_agent(self.llm, self.tools)

    def transcribe_audio(self, audio_filepath):
        """Transcrit un fichier audio en texte."""
        print(f"Tentative de transcription du fichier audio : {audio_filepath}")
        if audio_filepath is None:
            print("Chemin du fichier audio est None. Retour de 'Aucun enregistrement.'")
            return "Aucun enregistrement."

        recognizer = sr.Recognizer()
        try:
            with sr.AudioFile(audio_filepath) as source:
                print("Lecture du fichier audio...")
                audio = recognizer.record(source)  # lire le fichier audio entier
            
            print("Transcription de l'audio...")
            # Utiliser la reconnaissance Google pour le français
            text = recognizer.recognize_google(audio, language='fr-FR')
            print(f"Transcription réussie : {text}")
            # Appliquer la correction de transcription si nécessaire
            corrected_text = self.correct_transcription(text)
            print(f"Transcription corrigée : {corrected_text}")
            return corrected_text
        except sr.UnknownValueError:
            print("La reconnaissance vocale n'a pas pu comprendre l'audio")
            return "La reconnaissance vocale n'a pas compris l'audio."
        except sr.RequestError as e:
            print(f"Impossible d'obtenir les résultats du service de reconnaissance vocale Google ; {e}")
            return f"Erreur de service de reconnaissance vocale ; {e}"
        except Exception as e:
             print(f"Une erreur inattendue s'est produite lors de la transcription : {e}")
             return f"Une erreur inattendue s'est produite lors de la transcription : {e}"

    def speak(self, text):
        """Convertit le texte en parole en utilisant gTTS et retourne le chemin du fichier audio."""
        if not text:
            # Rien à dire si le texte est vide
            return None
        try:
            tts = gTTS(text=text, lang='fr')
            # Utiliser un fichier temporaire pour sauvegarder l'audio généré
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
            temp_file.close()
            tts.save(temp_file.name)
            # Retourner le chemin du fichier temporaire
            return temp_file.name
        except Exception as e:
            print(f"Erreur lors de la synthèse vocale : {e}")
            return None # Retourner None en cas d'erreur

    def correct_transcription(self, text):
        """Corrige les erreurs de transcription tout en préservant le sens."""
        # Ne pas essayer de corriger les messages d'erreur ou les non-transcriptions
        if text in ["La reconnaissance vocale n'a pas compris l'audio.", "Erreur de service de reconnaissance vocale", "Aucun enregistrement.", "Une erreur inattendue s'est produite lors de la transcription"]:
            return text

        print("Tentative de correction de transcription...")
        prompt = f"""
        Corrigez les erreurs de transcription évidentes dans le texte suivant tout en préservant strictement le sens, le style et le contenu original.
        Ignorez les bruits de fond transcrits ou les hésitations si possible.
        Contexte (question précédente posée par le recruteur) : {self.current_question}
        
        Texte à corriger : {text}
        
        Fournissez uniquement la version corrigée, sans explications.
        """
        
        try:
            # Utiliser le modèle Gemini pour la correction
            correction_response = self.model.generate_content(prompt)
            # Vérifier si la réponse contient du texte avant de la retourner
            corrected_text = correction_response.text.strip() if correction_response and correction_response.text else text
            print(f"Correction de transcription terminée : {corrected_text}")
            return corrected_text
        except Exception as e:
            print(f"Erreur lors de la correction de transcription : {e}")
            return text # Retourner le texte original en cas d'erreur

    def extract_text_from_pdf(self, file):
        # Docstring déjà en français
        # print(f"Attempting to extract text from file: {file}") # Commenté car le paramètre file n'est pas un chemin ici, mais un objet _io.BufferedReader
        if file is None:
            print("Fichier fourni est None. Retour de chaîne vide.")
            return ""
        try:
            # Vérifier l'extension du fichier
            if file.name.lower().endswith('.pdf'):
                print("Extraction de texte à partir d'un fichier PDF...")
                reader = PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                print("Extraction PDF terminée.")
                return text
            elif file.name.lower().endswith('.docx'):
                print("Extraction de texte à partir d'un fichier DOCX...")
                doc = Document(file)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                print("Extraction DOCX terminée.")
                return text
            else:
                print("Format de fichier non supporté.")
            return "Format de fichier non supporté"
        except Exception as e:
             print(f"Erreur lors de l'extraction de texte : {e}")
             return "Erreur lors de la lecture du fichier."

    def analyze_cv(self, cv_file, job_description):
        """Analyse le CV et la description du poste en utilisant l'agent LangChain, retourne une analyse formatée et le texte brut du CV."""
        self.cv_text = ""
        self.job_description = job_description
        self.cv_analysis = None # Réinitialiser l'analyse au début

        if cv_file is None:
            return "Veuillez télécharger un CV.", "", None
            
        try:
            # Assurez-vous que cv_file est un chemin de fichier accessible par l'outil read_document_tool
            # Si Gradio fournit un objet fichier, vous devrez peut-être l'enregistrer temporairement
            # Pour l'instant, supposons que cv_file.name est le chemin
            file_path = cv_file.name
            print(f"Tentative de lecture du fichier CV : {file_path}")
            
            # Utiliser l'outil read_document_tool DIRECTEMENT pour lire le fichier
            cv_text = read_document_tool.invoke({"file_path": file_path})
            self.cv_text = cv_text

            if "Error" in self.cv_text or "Unsupported" in self.cv_text:
                 print(f"Erreur ou format non supporté lors de la lecture du CV : {self.cv_text}")
                 return self.cv_text, "", None # Retourner le message d'erreur si la lecture échoue et pas d'analyse
            
            print("Lecture du CV réussie. Texte extrait.")

        except Exception as e:
            print(f"Erreur lors de la lecture du fichier CV avec l'outil : {e}")
            return f"Erreur lors de la lecture du fichier CV ou de l'appel de l'outil de lecture : {e}", "", None # Retourner None pour l'analyse

        if not self.cv_text:
             print("Le texte extrait du CV est vide.")
             return "Impossible d'extraire le texte du CV.", "", None # Retourner None pour l'analyse
             
        try:
            print("Utilisation de l'outil d'analyse de CV pour analyser le texte...")
            # Appeler DIRECTEMENT l'outil analyze_cv_with_llm_tool
            analysis_text = analyze_cv_with_llm_tool.invoke({"job_description": job_description, "cv_text": self.cv_text})
            self.cv_analysis = analysis_text # Récupérer la sortie de l'outil
            
            print("Analyse du CV par l'outil terminée.")
            
            # Affichage de l'analyse dans la console
            self.secretaire.display_analysis(self.cv_analysis)
            
            # Retourner l'analyse et le texte du CV
            return self.cv_analysis, self.cv_text, self.cv_analysis # Retourner l'analyse pour l'affichage ET pour l'état
        except Exception as e:
            print(f"Erreur lors de l'analyse du CV par l'outil IA : {e}")
            self.cv_analysis = f"Erreur lors de l'analyse du CV par l'outil IA : {e}"
            # Retourner le message d'erreur pour l'analyse et le texte du CV, et None pour l'état de l'analyse
            return self.cv_analysis, self.cv_text, None

    def determine_interview_goals(self):
        """Détermine les objectifs de l'entretien basés sur le CV et la description du poste"""
        prompt = f"""
        Basé sur cette analyse de CV:
        {self.cv_analysis}
        
        Identifie 3 à 5 objectifs clés que l'entretien doit atteindre pour évaluer correctement le candidat.
        Ces objectifs doivent être spécifiques et axés sur l'évaluation des compétences et de l'adéquation au poste.
        Format: Liste simple d'objectifs, un par ligne.
        Exemples: Évaluer l'expérience en gestion de projet, Vérifier les compétences techniques en Python, Comprendre la motivation pour le poste.
        """
        
        try:
            response = self.model.generate_content(prompt)
            # Filtrer les lignes vides et s'assurer qu'il y a des objectifs
            goals = [goal.strip() for goal in response.text.split('\n') if goal.strip()]
            self.interview_goals = goals if goals else ["Évaluer l'adéquation générale du candidat."]
            return self.interview_goals
        except Exception as e:
            print(f"Erreur lors de la détermination des objectifs de l'entretien : {e}")
            self.interview_goals = ["Impossible de déterminer les objectifs. Évaluation générale."]
            return self.interview_goals


    def evaluate_goal_achievement(self, question, response):
        """Évalue si un objectif a été atteint par la question/réponse actuelle."""
        # Cette logique d'évaluation doit être plus fine. Au lieu de simplement vérifier si *un* objectif
        # a été atteint, l'IA devrait identifier *quel(s)* objectif(s) ont été potentiellement couverts.
        # Cependant, pour une implémentation simplifiée avec l'outil generate_content, nous allons
        # demander à l'IA si cette interaction a permis de progresser significativement vers un ou plusieurs objectifs.
        
        if not self.interview_goals:
            return False # Pas d'objectifs définis

        prompt = f"""
        En se basant sur la question et la réponse suivantes, et les objectifs de l'entretien:
        
        Objectifs de l'entretien: {self.interview_goals}
        Objectifs déjà considérés comme partiellement ou totalement atteints (basé sur les questions précédentes): {list(self.achieved_goals)}
        
        Question actuelle: {question}
        Réponse du candidat: {response}
        
        Cette interaction (question + réponse) a-t-elle significativement contribué à évaluer le candidat sur un ou plusieurs des objectifs NON ENCORE ATTEINTS ?
        Réponds uniquement par "OUI" ou "NON".
        """
        
        try:
            evaluation_response = self.model.generate_content(prompt)
            # Vérifier si la réponse est valide avant de la comparer
            if evaluation_response and evaluation_response.text:
                return evaluation_response.text.strip().upper() == "OUI"
            return False # Considérer non atteint si l'évaluation échoue
        except Exception as e:
            print(f"Erreur lors de l'évaluation de l'objectif : {e}")
            return False # En cas d'erreur, ne pas considérer l'objectif comme atteint


    def should_continue_interview(self):
        """Détermine si l'entretien doit continuer en se basant sur les objectifs et la progression."""
        # L'entretien devrait continuer si tous les objectifs n'ont pas été évalués de manière satisfaisante
        # ET si l'IA pense qu'il est encore possible d'obtenir des informations utiles.

        if not self.interview_goals:
             # Si pas d'objectifs, se baser sur un nombre max de questions ou un signal de l'IA
             return self.question_count < 5 # Limite arbitraire si pas d'objectifs clairs
             
        # Si tous les objectifs considérés comme atteints (simplifié ici) et qu'il y a des objectifs définis
        if len(self.achieved_goals) >= len(self.interview_goals) and len(self.interview_goals) > 0:
             return False
             
        if self.question_count >= 15: # Limite maximale de questions pour éviter une boucle infinie
             return False
             
        # Demander à l'IA si continuer l'entretien est pertinent pour atteindre les objectifs restants
        prompt = f"""
        Objectifs de l'entretien: {self.interview_goals}
        Objectifs considérés comme atteints (basé sur l'évaluation simplifiée): {list(self.achieved_goals)}
        Transcription actuelle de l'entretien:
        {self.interview_transcript}

        Compte tenu de la progression et des informations obtenues, est-il judicieux de poser une autre question pour mieux évaluer le candidat sur les objectifs non encore couverts?
        Réponds uniquement par "OUI" ou "NON".
        """

        try:
            continuation_response = self.model.generate_content(prompt)
            if continuation_response and continuation_response.text:
                 return continuation_response.text.strip().upper() == "OUI"
            return True # Continuer par défaut si l'évaluation de continuation échoue
        except Exception as e:
            print(f"Erreur lors de l'évaluation de la continuation : {e}")
            return True # Continuer par défaut en cas d'erreur


    def generate_next_question(self):
        """Génère la prochaine question basée sur le contexte de l'entretien et les objectifs restants."""
        print("Début de la génération de la prochaine question...")
        print(f"CV Analysis: {self.cv_analysis}")
        print(f"Interview Goals: {self.interview_goals}")
        print(f"Achieved Goals: {self.achieved_goals}")
        print(f"Interview Transcript: {self.interview_transcript}")

        prompt = f"""
        Tu es un recruteur professionnel qui mène un entretien d'embauche.
        Basé sur cette analyse de CV et l'historique de l'entretien:
        {self.cv_analysis}
        
        Objectifs de l'entretien: {self.interview_goals}
        Objectifs considérés comme atteints: {list(self.achieved_goals)}
        Questions et réponses précédentes:
        {self.interview_transcript}
        
        En gardant à l'esprit les objectifs NON ENCORE SUFFISAMMENT COUVERTS, formule une question d'entretien pertinente et ouverte pour continuer d'évaluer le candidat.
        La question doit être directe, professionnelle et s'appuyer sur la conversation précédente.
        IMPORTANT: Réponds UNIQUEMENT avec la question, sans explications ni commentaires.
        """
        
        try:
            print("Envoi de la requête à l'API Gemini...")
            response = self.model.generate_content(prompt)
            print(f"Réponse reçue de l'API: {response.text if response and response.text else 'Pas de réponse'}")
            
            if not response or not response.text:
                print("Pas de réponse de l'API, utilisation d'une question par défaut")
                return "Pouvez-vous me parler de votre expérience professionnelle la plus significative ?"
                
            return response.text.strip()
        except Exception as e:
            print(f"Erreur lors de la génération de la question suivante : {e}")
            return "Pouvez-vous me parler de votre expérience professionnelle la plus significative ?"

    def format_conversation_for_display(self, history):
        """Formate l'historique de conversation (liste de dict) en une chaîne Markdown pour l'affichage."""
        formatted_text = ""
        for message in history:
            role = message.get('role', 'Inconnu')
            content = message.get('content', '')
            
            if role == 'Recruteur':
                # Pour le recruteur, le contenu est juste la question texte
                formatted_text += f"**Recruteur :** {content}\n\n"
            elif role == 'Candidat':
                # Pour le candidat, le contenu est la réponse transcrite
                formatted_text += f"**Candidat :** {content}\n\n"
            # Ignorer les autres rôles ou formats si nécessaire
            
        return formatted_text

    def start_interview(self, cv_analysis_state):
        """Démarre l'entretien en utilisant l'agent LangChain pour la première question."""
        print("\n--- Démarrage de l'entretien (Mode LangChain Agent) ---")
        # Utiliser la valeur de l'état Gradio pour la vérification
        if not cv_analysis_state:
            print("Pas d'analyse de CV disponible dans l'état Gradio. Impossible de démarrer l'entretien.")
            # Retourner les états pour désactiver les contrôles et afficher un message
            return [], "", "Veuillez d'abord analyser un CV.", "", gr.update(interactive=False), gr.update(interactive=False)
        
        self.interview_started = True
        self.question_count = 0
        self.interview_transcript = []
        self.achieved_goals = set()
        self.conversation_history = [] # Réinitialiser l'historique au début
        print("État de l'entretien initialisé.")
        
        print("Détermination des objectifs de l'entretien...")
        # La détermination des objectifs peut aussi être un rôle pour l'agent Interviewer ou un outil spécifique.
        # Pour l'instant, gardons-la séparée ou intégrons-la comme une étape initiale de l'agent.
        # self.determine_interview_goals() # On réutilise la méthode existante pour l'instant
        
        # Utiliser l'agent Interviewer pour déterminer les objectifs via le nouvel outil
        try:
            print("Appel de l'outil determine_interview_goals_tool pour déterminer les objectifs...")
            # Appeler DIRECTEMENT l'outil determine_interview_goals_tool
            goals_list = determine_interview_goals_tool.invoke({
                "cv_analysis": cv_analysis_state, # Utiliser l'analyse du CV de l'état
                "job_description": self.job_description # Utiliser la description du poste stockée
            })
            
            self.interview_goals = goals_list
            print(f"Objectifs déterminés par l'outil : {self.interview_goals}")
            
            if not isinstance(self.interview_goals, list) or not self.interview_goals:
                 print("Outil n'a pas retourné une liste d'objectifs valide ou la liste est vide. Utilisation des objectifs par défaut.")
                 self.interview_goals = ["Impossible de déterminer les objectifs. Évaluation générale."]

        except Exception as e:
            print(f"Erreur lors de l'appel de l'outil determine_interview_goals_tool : {e}")
            self.interview_goals = [f"Erreur lors de la détermination des objectifs : {e}"] # Objectifs d'erreur

        # Assurer que les objectifs sont formatés avec des tirets pour l'affichage
        cleaned_goals = [goal.lstrip('* -').strip() for goal in self.interview_goals]
        goals_text = "Objectifs :\n" + "\n".join([f"- {goal}" for goal in cleaned_goals])
        print(f"Objectifs déterminés (nettoyés et formatés): \n{goals_text}")
        
        print("Génération de la première question via l'agent LangChain...")
        # L'agent doit générer la première question sans réponse précédente
        try:
            # L'agent décide quoi faire en fonction du prompt initial et des outils.
            # L'input est le contexte pour l'agent au début de l'entretien.
            agent_response = self.interview_agent.invoke({
                "cv_analysis": self.cv_analysis,
                "interview_goals": self.interview_goals,
                "interview_transcript": self.interview_transcript, # Vide au début
                "candidate_response": "" # Vide au début
            })
            question = agent_response['output'].strip()
            print(f"Première question générée par l'agent: '{question}'")
            self.current_question = question
            
            if not question:
                print("Agent n'a pas généré de question. Utilisation par défaut.")
                question = "Pouvez-vous me parler de votre expérience professionnelle la plus significative ?"
                self.current_question = question

        except Exception as e:
            print(f"Erreur lors de l'appel de l'agent Interviewer pour la première question : {e}")
            question = "Pouvez-vous me parler de votre expérience professionnelle la plus significative ?" # Question par défaut en cas d'erreur
            self.current_question = question

        # Le contenu pour l'historique interne est simplement la question texte
        chat_message_content = question
        
        # Ajouter le premier message à l'historique interne
        first_message_for_history = {'role': 'Recruteur', 'content': chat_message_content}
        self.conversation_history.append(first_message_for_history)
        print(f"Premier message ajouté à l'historique interne: {self.conversation_history}")

        # --- Nouveau : Formater l'historique interne pour l'affichage ---
        formatted_conversation = self.format_conversation_for_display(self.conversation_history)
        print(f"Conversation formatée pour l'affichage initial: \n{formatted_conversation}")
        # --- Fin Nouveau ---

        # Ajouter un petit délai pour permettre à Gradio de traiter la mise à jour de l'état
        time.sleep(0.5)

        print("\n--- Fin de start_interview ---")

        # Retourner l'historique interne et la chaîne formatée pour l'affichage, plus les autres outputs
        # Outputs: [conversation_history_state, conversation_display, goals_display, response_input, audio_input, submit_btn]
        # On active l'input audio et on désactive le bouton submit initialement
        return self.conversation_history, formatted_conversation, goals_text, "", gr.update(interactive=True), gr.update(interactive=False)

    def get_recorded_response(self, audio_file):
         """Utilise le fichier audio enregistré par Gradio pour transcrire la réponse."""
         print(f"Fichier audio reçu de Gradio : {audio_file}")
         if audio_file is None:
             print("Aucun fichier audio reçu.")
             # Retourner une chaîne vide, désactiver le bouton soumettre et None pour le chemin du fichier
             return "Aucun enregistrement fourni.", gr.update(interactive=False), None
             
         # Stocker le chemin du fichier audio dans l'état pour une utilisation ultérieure
         # Le chemin est nécessaire pour process_candidate_response
         audio_filepath_for_state = audio_file # Stocker le chemin reçu

         response_text = self.transcribe_audio(audio_file)
         # Après la transcription, le bouton Soumettre doit être actif (il sera géré dans process_candidate_response)
         # Le bouton Enregistrer redevient interactif ici si nécessaire
         return response_text, gr.update(interactive=True), audio_filepath_for_state

    def process_candidate_response(self, conversation_history, audio_file):
        """Traite la réponse du candidat (via fichier audio) en utilisant l'agent LangChain."""
        print("\n--- Traitement de la réponse du candidat par l'agent LangChain ---")

        # Assurez-vous que conversation_history est une liste
        if not isinstance(conversation_history, list):
            print(f"Attention: conversation_history n'est pas une liste, type reçu: {type(conversation_history)}. Initialisation à une liste vide.")
            conversation_history = []
            self.conversation_history = conversation_history[:] # Synchroniser l'historique interne
        else:
            # Utiliser l'historique passé en input qui est géré par Gradio State
            self.conversation_history = conversation_history[:] # Synchroniser l'historique interne

        if not self.interview_started or self.current_question is None:
            print("Entretien non démarré ou question actuelle manquante. Arrêt du traitement.")
            formatted_conversation = self.format_conversation_for_display(self.conversation_history)
            # Retourner les sorties nécessaires, désactiver les contrôles
            return self.conversation_history, formatted_conversation, gr.update(value=""), gr.update(value=None, interactive=False), gr.update(interactive=False)

        # L'audio a déjà été transcrit dans get_recorded_response, audio_file est le chemin stocké
        # Nous devons récupérer la réponse transcrite. Elle est déjà dans response_input.
        # Nous allons donc simplement utiliser le chemin du fichier ici pour l'analyse sentiment/ajout à la transcription si nécessaire
        # La transcription réelle est gérée par get_recorded_response et mise à jour dans response_input.

        # Nous devons récupérer la transcription qui a été mise à jour par get_recorded_response dans response_input
        # Cependant, la logique actuelle ne passe pas response_input ici. Il faut la passer comme input.
        # Pour l'instant, je vais refaire la transcription ici en utilisant le chemin du fichier.
        # TODO: Passer response_input comme input à cette fonction pour éviter la double transcription.

        # 1. Transcrire l'audio (ou utiliser la transcription déjà faite si disponible et passée en input)
        # Pour l'instant, refaisons la transcription car response_input n'est pas un input ici
        response = self.transcribe_audio(audio_file)
        print(f"Réponse du candidat transcrite (dans process_candidate_response): '{response}'")

        # 2. Mettre à jour l'historique interne et la transcription détaillée
        candidate_message_for_history = {'role': 'Candidat', 'content': response}
        self.conversation_history.append(candidate_message_for_history)
        print(f"Réponse du candidat ajoutée à l'historique interne. Historique interne actuel: {self.conversation_history}")
        
        # Ajouter la question précédente, la réponse actuelle et l'analyse à la transcription
        if self.current_question: # S'assurer qu'une question a été posée
            # Effectuer l'analyse de sentiment ici pour la transcription détaillée
            sentiment = self.analyze_sentiment(response)
            self.interview_transcript.append({
                'question': self.current_question,
                'reponse': response,
                'sentiment': sentiment # Ajouter le sentiment analysé
            })
            print(f"Transcription de l'entretien détaillée mise à jour: {self.interview_transcript}")
        else:
            print("Avertissement: self.current_question est None lors de l'ajout à la transcription détaillée.")

        self.question_count += 1
        print(f"Compteur de questions: {self.question_count}")
        
        # 3. Appeler l'agent Simulateur d'Entretien avec le contexte mis à jour
        print("Appel de l'agent Simulateur d'Entretien pour la prochaine étape...")
        try:
            agent_response = self.interview_agent.invoke({
                "cv_analysis": self.cv_analysis,
                "interview_goals": self.interview_goals,
                "interview_transcript": self.interview_transcript,
                "candidate_response": response # Passer la réponse du candidat à l'agent
            })
            agent_output = agent_response['output'].strip()
            print(f"Sortie brute de l'agent Interviewer: '{agent_output}'")

            # Validation et interprétation de la sortie de l'agent
            if not agent_output or agent_output.upper() == "FIN":
                print("Agent a signalé la fin de l'entretien ou a retourné une sortie vide.")
                self.interview_started = False
                final_message = "L'entretien est terminé. Vous pouvez maintenant générer le rapport final." # Message de fin
                final_message_for_history = {'role': 'Recruteur', 'content': final_message}
                self.conversation_history.append(final_message_for_history)
                print(f"Message de fin ajouté à l'historique interne. Historique interne actuel: {self.conversation_history}")
                
                formatted_conversation = self.format_conversation_for_display(self.conversation_history)

                # Désactiver les contrôles d'entretien et retourner les états finaux
                return (self.conversation_history, 
                        formatted_conversation, 
                        gr.update(value="", interactive=False), # Désactiver la zone de réponse
                        gr.update(value=None, interactive=False), # Désactiver l'enregistrement audio
                        gr.update(interactive=False)) # Désactiver le bouton soumettre

            # L'agent a généré une nouvelle question
            next_question = agent_output
            print(f"Prochaine question déterminée : '{next_question}'")
            self.current_question = next_question # Mettre à jour la question actuelle

            # Ajouter la question suivante à l'historique interne
            next_question_for_history = {'role': 'Recruteur', 'content': next_question}
            self.conversation_history.append(next_question_for_history)
            print(f"Prochaine question ajoutée à l'historique interne. Historique interne actuel: {self.conversation_history}")
            
            # Mettre à jour l'affichage et réactiver les contrôles pour le prochain tour
            formatted_conversation = self.format_conversation_for_display(self.conversation_history)
            print(f"Conversation formatée pour l'affichage : \n{formatted_conversation}")

            # Retourner les états mis à jour pour le prochain tour de conversation
            return (self.conversation_history, 
                    formatted_conversation, 
                    gr.update(value="", interactive=True), # Réactiver la zone de réponse
                    gr.update(value=None, interactive=True), # Réactiver l'enregistrement audio
                    gr.update(interactive=False)) # Désactiver le bouton soumettre jusqu'à nouvel enregistrement/saisie

        except Exception as e:
            print(f"Erreur lors de l'appel de l'agent Interviewer pour traiter la réponse : {e}")
            error_message = f"Une erreur est survenue lors de la génération de la prochaine question : {e}. L'entretien est terminé."
            self.interview_started = False # Marquer la fin de l'entretien en cas d'erreur
            error_message_for_history = {'role': 'Recruteur', 'content': error_message}
            self.conversation_history.append(error_message_for_history)
            formatted_conversation = self.format_conversation_for_display(self.conversation_history)
            # Retourner les états finaux, désactiver les contrôles
            return self.conversation_history, formatted_conversation, gr.update(value="", interactive=False), gr.update(value=None, interactive=False), gr.update(interactive=False)

    def analyze_sentiment(self, text):
        """Analyse le sentiment du texte donné en utilisant le modèle Gemini via LangChain."""
        # Vérifier les messages spéciaux qui ne doivent pas être analysés pour le sentiment
        if text in ["La reconnaissance vocale n'a pas compris l'audio.", "Erreur de service de reconnaissance vocale", "Aucun enregistrement.", "Une erreur inattendue s'est produite lors de la transcription"]:
            return "Non analysable"

        prompt = f"""
        Analyse le sentiment principal de la réponse suivante du candidat dans le contexte d'un entretien d'embauche. Est-il positif, négatif, neutre, hésitant, confiant, stressé, enthousiaste ? Choisis le terme qui décrit le mieux le sentiment général.
        
        Réponse du candidat: {text}
        
        Réponds uniquement avec un seul terme ou une courte phrase décrivant le sentiment.
        """
        try:
            # Utiliser self.llm (le modèle LangChain) pour l'analyse de sentiment
            response = self.llm.invoke(prompt)
            # S'assurer qu'il y a du texte dans la réponse avant de la strip
            sentiment = response.content.strip() if response and response.content else "Analyse impossible"
            print(f"Analyse de sentiment via LLM terminée : {sentiment}") # Message de debug
            return sentiment
        except Exception as e:
            print(f"Erreur lors de l'analyse de sentiment via LLM : {e}")
            return f"Erreur lors de l'analyse de sentiment : {e}"

    def save_analysis(self, analysis_text):
        """Sauvegarde l'analyse dans un fichier PDF professionnel."""
        print("Attempting to save analysis to PDF...") # Message de debug
        if not analysis_text or analysis_text.startswith("Erreur") or analysis_text.startswith("Impossible"):
            print("Pas d'analyse valide à sauvegarder.") # Message de debug
            return None # Retourner None si l'analyse est vide ou contient une erreur
        try:
            # Créer un fichier PDF temporaire
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            temp_file.close()
            pdf_path = temp_file.name
            print(f"Fichier temporaire créé pour l'analyse PDF : {pdf_path}") # Message de debug
            
            # Utiliser l'agent secrétaire pour créer le PDF
            self.secretaire.create_pdf_report(analysis_text, pdf_path)
            print(f"PDF d'analyse créé à : {pdf_path}") # Message de debug
            # Vérifier si le fichier a été créé et n'est pas vide
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                 return pdf_path
            else:
                 print("Erreur : Fichier PDF d'analyse créé est vide ou inexistant.") # Message de debug
            return None
        except Exception as e:
            print(f"Erreur lors de la sauvegarde de l'analyse en PDF : {e}") # Message de debug
            return None # Retourner None en cas d'erreur
            
    def generate_report(self):
        """Génère le rapport final de l'entretien."""
        print("Attempting to generate final report...") # Message de debug
        if not self.cv_analysis or not self.interview_transcript or not self.interview_goals:
            print("Données insuffisantes pour le rapport final.") # Message de debug
            # Retourner None car le composant File attend un chemin ou None
            return None
            
        try:
            print("Appel de l'outil generate_final_report_content_tool...") # Message de debug
            # Appeler DIRECTEMENT l'outil generate_final_report_content_tool
            evaluation_json_string = generate_final_report_content_tool.invoke({
                "cv_analysis": self.cv_analysis,
                "interview_transcript": self.interview_transcript,
                "interview_goals": self.interview_goals
            })
            print("Outil generate_final_report_content_tool terminé.") # Message de debug

            # S'assurer que l'évaluation est une chaîne JSON valide
            print(f"Sortie brute de l'outil : {evaluation_json_string[:500]}...") # Message de debug

            # create_final_report_pdf a déjà une logique pour gérer le JSON potentiellement invalide en retournant un rapport par défaut

            # Créer le rapport PDF avec l'évaluation JSON
            print("Création du rapport PDF final...") # Message de debug
            pdf_path = self.secretaire.create_final_report_pdf(
                self.cv_analysis,
                self.interview_transcript,
                evaluation_json_string # Passer la chaîne JSON (même si potentiellement invalide, gérée par create_final_report_pdf)
            )

            if pdf_path and os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                print(f"Rapport PDF final généré avec succès : {pdf_path}") # Message de debug
                return pdf_path
            else:
                print("Erreur lors de la création du fichier PDF final ou fichier vide.") # Message de debug
                return None # Retourner None si la création échoue ou fichier vide

        except Exception as e:
            print(f"Erreur inattendue lors de la génération du rapport final : {e}") # Message de debug
            return None # Retourner None en cas d'erreur

def create_interface():
    app = RecrutementApp(llm, tools)
    
    with gr.Blocks(title="Tomouhi") as interface:
        gr.Markdown("Toumouhi")
        
        with gr.Tab("Analyse de CV"):
            cv_analysis_state = gr.State(None)
            
            with gr.Column():
                cv_file = gr.File(label="Télécharger le CV (PDF ou DOCX)")
                job_desc = gr.Textbox(label="Description du poste", lines=5)
                analyze_btn = gr.Button("Analyser le CV")
                
                cv_analysis_display = gr.Textbox(label="Analyse du CV", lines=15, interactive=False, render="markdown")
                
                download_analysis_btn = gr.Button("Télécharger l'analyse", interactive=False)
                analysis_download_file = gr.File(label="Fichier d'analyse", interactive=False)
            
            analyze_btn.click(
                app.analyze_cv,
                inputs=[cv_file, job_desc],
                outputs=[cv_analysis_display, gr.State(app.cv_text), cv_analysis_state]
            ).success(
                lambda analysis_result: [gr.update(interactive=bool(analysis_result)), None] if analysis_result else [gr.update(interactive=False), None],
                inputs=cv_analysis_display,
                outputs=[download_analysis_btn, analysis_download_file]
            )
            
            download_analysis_btn.click(
                app.save_analysis,
                inputs=cv_analysis_display,
                outputs=analysis_download_file
            )
            
        
        with gr.Tab("Simulation d'Entretien"):
            conversation_history_state = gr.State([])
            conversation_display = gr.Markdown("", label="Conversation de l'entretien")
            
            last_audio_file_state = gr.State(None)
            
            goals_display = gr.Textbox(label="Objectifs de l'entretien", lines=5, interactive=False)
            start_interview_btn = gr.Button("Démarrer l'entretien")
            
            with gr.Row():
                response_input = gr.Textbox(label="Votre réponse", lines=3, scale=3, interactive=False)
                with gr.Column(scale=1):
                    audio_input = gr.Audio(sources="microphone", type="filepath", label="Enregistrer votre réponse")
                    submit_btn = gr.Button("Soumettre réponse", interactive=False)

            start_interview_event = start_interview_btn.click(
                app.start_interview,
                inputs=[cv_analysis_state],
                outputs=[
                    conversation_history_state,
                    conversation_display,
                    goals_display, 
                    response_input, 
                    audio_input,
                    submit_btn
                ]
            )

            audio_input.change(
                 app.get_recorded_response,
                 inputs=audio_input,
                 outputs=[response_input, submit_btn, last_audio_file_state]
            )
            
            submit_btn.click(
                app.process_candidate_response,
                inputs=[conversation_history_state, last_audio_file_state],
                outputs=[
                     conversation_history_state,
                     conversation_display,
                     response_input,
                     audio_input,
                     submit_btn
                ]
            )
            
        
        with gr.Tab("Rapport Final"):
            generate_report_btn = gr.Button("Générer le rapport")
            report_output = gr.File(label="Télécharger le rapport")
            
            generate_report_btn.click(
                app.generate_report,
                outputs=report_output
            )
    
    return interface

if __name__ == "__main__":
    interface = create_interface()
    # Pour le déploiement sur Hugging Face Spaces
    interface.launch() 